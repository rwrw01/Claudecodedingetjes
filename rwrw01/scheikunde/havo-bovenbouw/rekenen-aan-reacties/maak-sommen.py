"""Genereer twee Word-documenten: sommen + sommen+antwoorden voor HAVO 4."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OPGAVEN = [
    {
        "nr": 1,
        "titel": "Magnesium verbrandt in zuurstof (massa → massa, 2:2-verhouding)",
        "vraag": (
            "Bij het verbranden van magnesium in lucht ontstaat magnesiumoxide (MgO). "
            "Een leerling verbrandt 8,0 g magnesium volledig.\n\n"
            "Bereken hoeveel gram magnesiumoxide er ontstaat.\n\n"
            "Gegeven: M(Mg) = 24,3 g/mol ; M(MgO) = 40,3 g/mol"
        ),
        "antwoord_stappen": [
            "Stap 1 — Reactievergelijking kloppend maken:  2 Mg + O₂ → 2 MgO",
            "Stap 2 — Mol Mg berekenen:  n(Mg) = 8,0 / 24,3 = 0,329 mol",
            "Stap 3 — Molverhouding Mg : MgO = 2 : 2 = 1 : 1",
            "Stap 4 — Mol MgO:  n(MgO) = 0,329 × (2/2) = 0,329 mol",
            "Stap 5 — Massa MgO:  m(MgO) = 0,329 × 40,3 = 13,3 g",
        ],
        "antwoord": "Er ontstaat 13,3 g magnesiumoxide.",
    },
    {
        "nr": 2,
        "titel": "Aluminium met zuurstof (mol → mol, 4:3-verhouding)",
        "vraag": (
            "Aluminiumpoeder reageert met zuurstofgas tot aluminiumoxide (Al₂O₃). "
            "In een experiment wordt 0,60 mol aluminium volledig omgezet.\n\n"
            "a) Stel de kloppende reactievergelijking op.\n"
            "b) Bereken hoeveel mol zuurstofgas (O₂) nodig is.\n"
            "c) Bereken hoeveel mol aluminiumoxide ontstaat."
        ),
        "antwoord_stappen": [
            "Stap 1 — Reactievergelijking:  4 Al + 3 O₂ → 2 Al₂O₃",
            "Stap 2 — Molverhouding Al : O₂ : Al₂O₃ = 4 : 3 : 2",
            "Stap 3 — n(O₂) = 0,60 × (3/4) = 0,45 mol",
            "Stap 4 — n(Al₂O₃) = 0,60 × (2/4) = 0,30 mol",
        ],
        "antwoord": "b) 0,45 mol O₂   c) 0,30 mol Al₂O₃",
    },
    {
        "nr": 3,
        "titel": "IJzer + zwavel (overmaat berekenen)",
        "vraag": (
            "In een proef worden 7,00 g ijzer en 5,00 g zwavel samen verhit. "
            "Ze reageren tot ijzer(II)sulfide (FeS).\n\n"
            "a) Stel de reactievergelijking op.\n"
            "b) Bepaal welke stof in overmaat is.\n"
            "c) Bereken hoeveel gram FeS er ontstaat.\n"
            "d) Bereken hoeveel gram van de overmaat-stof er overblijft.\n\n"
            "Gegeven: M(Fe) = 55,8 g/mol ; M(S) = 32,1 g/mol ; M(FeS) = 87,9 g/mol"
        ),
        "antwoord_stappen": [
            "Stap 1 — Reactievergelijking:  Fe + S → FeS  (verhouding 1 : 1 : 1)",
            "Stap 2 — Mol berekenen:",
            "    n(Fe) = 7,00 / 55,8 = 0,1254 mol",
            "    n(S)  = 5,00 / 32,1 = 0,1558 mol",
            "Stap 3 — Vergelijk met molverhouding 1:1. Er is minder Fe dan S, "
            "dus Fe is de beperkende stof en S is in overmaat.",
            "Stap 4 — n(FeS) = n(Fe) = 0,1254 mol",
            "    m(FeS) = 0,1254 × 87,9 = 11,0 g",
            "Stap 5 — Zwavel dat reageert: 0,1254 mol",
            "    Overgebleven S: 0,1558 − 0,1254 = 0,0304 mol",
            "    m(S over) = 0,0304 × 32,1 = 0,98 g",
        ],
        "antwoord": "b) S is in overmaat   c) 11,0 g FeS   d) 0,98 g S blijft over",
    },
    {
        "nr": 4,
        "titel": "Verbranding methaan (limiting reagent + massa)",
        "vraag": (
            "Bij de volledige verbranding van methaan (CH₄) ontstaan koolstofdioxide en water.\n\n"
            "Een gasbrander krijgt 16,0 g methaan en 48,0 g zuurstof toegevoerd.\n\n"
            "a) Stel de kloppende reactievergelijking op.\n"
            "b) Welke stof is beperkend, methaan of zuurstof?\n"
            "c) Bereken hoeveel gram koolstofdioxide er kan ontstaan.\n\n"
            "Gegeven: M(CH₄) = 16,0 g/mol ; M(O₂) = 32,0 g/mol ; M(CO₂) = 44,0 g/mol"
        ),
        "antwoord_stappen": [
            "Stap 1 — Reactievergelijking:  CH₄ + 2 O₂ → CO₂ + 2 H₂O",
            "Stap 2 — Mol berekenen:",
            "    n(CH₄) = 16,0 / 16,0 = 1,00 mol",
            "    n(O₂)  = 48,0 / 32,0 = 1,50 mol",
            "Stap 3 — Voor 1,00 mol CH₄ is 2,00 mol O₂ nodig, maar er is slechts 1,50 mol O₂.",
            "    Dus O₂ is de beperkende stof, CH₄ is in overmaat.",
            "Stap 4 — Molverhouding O₂ : CO₂ = 2 : 1",
            "    n(CO₂) = 1,50 × (1/2) = 0,75 mol",
            "Stap 5 — m(CO₂) = 0,75 × 44,0 = 33,0 g",
        ],
        "antwoord": "b) zuurstof (O₂) is beperkend   c) 33,0 g CO₂",
    },
    {
        "nr": 5,
        "titel": "Ammoniak verbranden (mol → massa, 4:6-verhouding)",
        "vraag": (
            "In de industrie wordt ammoniak (NH₃) verbrand met zuurstof tot stikstofmonoxide (NO) "
            "en water. Dit is de eerste stap bij de productie van salpeterzuur.\n\n"
            "Reactievergelijking: 4 NH₃ + 5 O₂ → 4 NO + 6 H₂O\n\n"
            "Bij een experiment reageert 0,80 mol NH₃ volledig.\n\n"
            "a) Bereken hoeveel mol zuurstof er minimaal nodig is.\n"
            "b) Bereken hoeveel gram water er ontstaat.\n\n"
            "Gegeven: M(H₂O) = 18,0 g/mol"
        ),
        "antwoord_stappen": [
            "Stap 1 — Reactievergelijking is gegeven en al kloppend.",
            "Stap 2 — Molverhouding NH₃ : O₂ : NO : H₂O = 4 : 5 : 4 : 6",
            "Stap 3 — n(O₂) = 0,80 × (5/4) = 1,00 mol",
            "Stap 4 — n(H₂O) = 0,80 × (6/4) = 1,20 mol",
            "Stap 5 — m(H₂O) = 1,20 × 18,0 = 21,6 g",
        ],
        "antwoord": "a) 1,00 mol O₂   b) 21,6 g water",
    },
    {
        "nr": 6,
        "titel": "Kalksteen branden voor cement (zelf de vergelijking opstellen)",
        "vraag": (
            "Bij de productie van cement wordt kalksteen sterk verhit. Kalksteen bestaat "
            "voornamelijk uit calciumcarbonaat (CaCO₃). Tijdens het branden ontleedt "
            "het in gebrande kalk (CaO) en koolstofdioxide.\n\n"
            "a) Stel zelf de kloppende reactievergelijking op.\n"
            "b) Bereken hoeveel kilogram CaO ontstaat uit 250 kg kalksteen "
            "(neem aan dat dit 100% CaCO₃ is).\n\n"
            "Gegeven: M(CaCO₃) = 100,1 g/mol ; M(CaO) = 56,1 g/mol ; M(CO₂) = 44,0 g/mol"
        ),
        "antwoord_stappen": [
            "Stap 1 — Reactievergelijking:  CaCO₃ → CaO + CO₂  (al kloppend, verhouding 1 : 1 : 1)",
            "Stap 2 — Massa naar gram:  250 kg = 250 000 g",
            "Stap 3 — n(CaCO₃) = 250 000 / 100,1 = 2 498 mol",
            "Stap 4 — Molverhouding CaCO₃ : CaO = 1 : 1, dus n(CaO) = 2 498 mol",
            "Stap 5 — m(CaO) = 2 498 × 56,1 = 140 138 g ≈ 140 kg",
        ],
        "antwoord": "b) Er ontstaat ongeveer 140 kg CaO (gebrande kalk).",
    },
    {
        "nr": 7,
        "titel": "Haber-Bosch ammoniaksynthese (beperkend reagens, vergelijking zelf opstellen)",
        "vraag": (
            "In de Haber-Bosch-fabriek wordt ammoniak (NH₃) gemaakt uit stikstofgas en "
            "waterstofgas. Een testreactor wordt gevuld met 56 g N₂ en 12 g H₂.\n\n"
            "a) Stel zelf de kloppende reactievergelijking op.\n"
            "b) Bepaal welke stof beperkend is.\n"
            "c) Bereken hoeveel gram ammoniak je maximaal kunt maken.\n\n"
            "Gegeven: M(N₂) = 28,0 g/mol ; M(H₂) = 2,02 g/mol ; M(NH₃) = 17,0 g/mol"
        ),
        "antwoord_stappen": [
            "Stap 1 — Reactievergelijking:  N₂ + 3 H₂ → 2 NH₃",
            "Stap 2 — Mol berekenen:",
            "    n(N₂) = 56 / 28,0 = 2,00 mol",
            "    n(H₂) = 12 / 2,02 = 5,94 mol",
            "Stap 3 — Voor 2,00 mol N₂ is volgens verhouding 1:3 nodig:  3 × 2,00 = 6,00 mol H₂.",
            "    Beschikbaar is 5,94 mol H₂ → H₂ is beperkend, N₂ is in overmaat.",
            "Stap 4 — Molverhouding H₂ : NH₃ = 3 : 2",
            "    n(NH₃) = 5,94 × (2/3) = 3,96 mol",
            "Stap 5 — m(NH₃) = 3,96 × 17,0 = 67,3 g",
        ],
        "antwoord": "b) H₂ is beperkend   c) maximaal 67,3 g NH₃",
    },
    {
        "nr": 8,
        "titel": "Propaan in een gas-bbq verbranden (vergelijking zelf balanceren)",
        "vraag": (
            "Propaan (C₃H₈) is het gas dat in een gas-bbq of campingfles zit. Bij "
            "volledige verbranding met zuurstof ontstaan koolstofdioxide en waterdamp.\n\n"
            "a) Stel zelf de kloppende reactievergelijking op.\n"
            "b) Bereken hoeveel gram CO₂ ontstaat als 0,40 mol propaan volledig verbrandt.\n"
            "c) Bereken hoeveel gram water er bij die reactie ontstaat.\n\n"
            "Gegeven: M(CO₂) = 44,0 g/mol ; M(H₂O) = 18,0 g/mol"
        ),
        "antwoord_stappen": [
            "Stap 1 — Balanceren:  C₃H₈ + 5 O₂ → 3 CO₂ + 4 H₂O",
            "    (controle: links 3 C, 8 H, 10 O ; rechts 3 C, 8 H, 10 O ✓)",
            "Stap 2 — Molverhouding C₃H₈ : O₂ : CO₂ : H₂O = 1 : 5 : 3 : 4",
            "Stap 3 — n(CO₂) = 0,40 × (3/1) = 1,20 mol",
            "    m(CO₂) = 1,20 × 44,0 = 52,8 g",
            "Stap 4 — n(H₂O) = 0,40 × (4/1) = 1,60 mol",
            "    m(H₂O) = 1,60 × 18,0 = 28,8 g",
        ],
        "antwoord": "b) 52,8 g CO₂   c) 28,8 g H₂O",
    },
    {
        "nr": 9,
        "titel": "Anorthiet + kalksteen — cementchemie (gevorderd, zelf balanceren)",
        "vraag": (
            "Anorthiet (CaAl₂Si₂O₈) is een mineraal dat in vulkanisch gesteente voorkomt. "
            "In de cementindustrie wordt het samen met kalksteen (CaCO₃) sterk verhit. "
            "Daarbij ontstaan een calcium-aluminosilicaat met formule Ca₃Al₂Si₂O₁₀ en "
            "koolstofdioxide.\n\n"
            "a) Stel zelf de kloppende reactievergelijking op. Tip: balanceer eerst "
            "Ca, dan Al en Si, en gebruik CO₃-eenheden voor de O- en C-balans.\n"
            "b) Bereken hoeveel kilogram kalksteen nodig is om 100 kg anorthiet "
            "volledig om te zetten.\n\n"
            "Gegeven: M(CaAl₂Si₂O₈) = 278,3 g/mol ; M(CaCO₃) = 100,1 g/mol"
        ),
        "antwoord_stappen": [
            "Stap 1 — Balanceren:  CaAl₂Si₂O₈ + 2 CaCO₃ → Ca₃Al₂Si₂O₁₀ + 2 CO₂",
            "    (controle: links 3 Ca, 2 Al, 2 Si, 14 O, 2 C ; rechts 3 Ca, 2 Al, 2 Si, "
            "10 + 4 = 14 O, 2 C ✓)",
            "Stap 2 — Massa naar gram:  100 kg = 100 000 g",
            "Stap 3 — n(CaAl₂Si₂O₈) = 100 000 / 278,3 = 359,3 mol",
            "Stap 4 — Molverhouding CaAl₂Si₂O₈ : CaCO₃ = 1 : 2",
            "    n(CaCO₃) = 359,3 × 2 = 718,7 mol",
            "Stap 5 — m(CaCO₃) = 718,7 × 100,1 = 71 942 g ≈ 71,9 kg",
        ],
        "antwoord": "b) Er is ongeveer 71,9 kg kalksteen (CaCO₃) nodig.",
    },
    {
        "nr": 10,
        "titel": "Zoutzuur op kalksteenvloer (beperkend reagens, zelf balanceren)",
        "vraag": (
            "Een schoonmaker giet per ongeluk zoutzuur op een kalksteenvloer. Het "
            "zoutzuur (HCl) reageert met de kalksteen (CaCO₃) waarbij calciumchloride "
            "(CaCl₂), water en koolstofdioxide ontstaan. Daardoor bruist het.\n\n"
            "Er komt 50 g kalksteen in contact met 30 g HCl.\n\n"
            "a) Stel zelf de kloppende reactievergelijking op.\n"
            "b) Bepaal welke stof beperkend is.\n"
            "c) Bereken hoeveel gram CO₂-gas er maximaal vrijkomt.\n\n"
            "Gegeven: M(CaCO₃) = 100,1 g/mol ; M(HCl) = 36,5 g/mol ; M(CO₂) = 44,0 g/mol"
        ),
        "antwoord_stappen": [
            "Stap 1 — Reactievergelijking:  CaCO₃ + 2 HCl → CaCl₂ + H₂O + CO₂",
            "Stap 2 — Mol berekenen:",
            "    n(CaCO₃) = 50 / 100,1 = 0,500 mol",
            "    n(HCl)   = 30 / 36,5  = 0,822 mol",
            "Stap 3 — Voor 0,500 mol CaCO₃ is volgens verhouding 1:2 nodig:  "
            "2 × 0,500 = 1,00 mol HCl. Beschikbaar 0,822 mol → HCl is beperkend.",
            "Stap 4 — Molverhouding HCl : CO₂ = 2 : 1",
            "    n(CO₂) = 0,822 / 2 = 0,411 mol",
            "Stap 5 — m(CO₂) = 0,411 × 44,0 = 18,1 g",
        ],
        "antwoord": "b) HCl is beperkend   c) maximaal 18,1 g CO₂",
    },
]


def maak_basis_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    titel = doc.add_heading('Sommen — Rekenen met reactievergelijkingen', level=0)
    titel.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('Scheikunde · HAVO 4 · 5 opgaven')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    intro = doc.add_paragraph()
    intro.add_run(
        'Werk elke som uit volgens het 5-stappenplan: reactievergelijking → mol → '
        'molverhouding → mol gevraagde stof → massa of mol als antwoord. '
        'Schrijf je tussenstappen op en let op significante cijfers (meestal 2 of 3).'
    ).italic = True

    doc.add_paragraph()

    return doc


def voeg_opgave_toe(doc, opgave, met_antwoord=False):
    h = doc.add_heading(f"Opgave {opgave['nr']} — {opgave['titel']}", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

    p = doc.add_paragraph(opgave['vraag'])
    p.paragraph_format.space_after = Pt(6)

    if met_antwoord:
        kop = doc.add_paragraph()
        kop_run = kop.add_run('Uitwerking')
        kop_run.bold = True
        kop_run.font.color.rgb = RGBColor(0xe1, 0x1d, 0x48)

        for stap in opgave['antwoord_stappen']:
            sp = doc.add_paragraph(stap)
            sp.paragraph_format.left_indent = Cm(0.5)
            sp.paragraph_format.space_after = Pt(2)

        ant = doc.add_paragraph()
        ant_run = ant.add_run(f"Antwoord: {opgave['antwoord']}")
        ant_run.bold = True
        ant_run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
    else:
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('Antwoord: ____________________________________________')

    doc.add_paragraph()


def main():
    out_dir = Path(__file__).parent

    doc_sommen = maak_basis_document()
    for opg in OPGAVEN:
        voeg_opgave_toe(doc_sommen, opg, met_antwoord=False)
    pad_sommen = out_dir / 'sommen-rekenen-aan-reacties.docx'
    doc_sommen.save(str(pad_sommen))
    print(f"Geschreven: {pad_sommen}")

    doc_ant = maak_basis_document()
    titel_p = doc_ant.paragraphs[0]
    titel_p.text = 'Sommen + antwoorden — Rekenen met reactievergelijkingen'
    titel_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titel_p.style = doc_ant.styles['Heading 1']
    for opg in OPGAVEN:
        voeg_opgave_toe(doc_ant, opg, met_antwoord=True)
    pad_ant = out_dir / 'sommen-rekenen-aan-reacties-antwoorden.docx'
    doc_ant.save(str(pad_ant))
    print(f"Geschreven: {pad_ant}")


if __name__ == '__main__':
    main()
